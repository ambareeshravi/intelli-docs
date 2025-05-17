import os
from typing import List, Dict, Any
from pathlib import Path
import PyPDF2
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document as LangchainDocument
from intelli_docs.core.config import settings

class DocumentProcessor:
    """
    Handles document processing and chunking
    These steps are required to interact with the LLMs in the appropriate way
    """
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.CHUNK_SIZE,
            chunk_overlap=settings.CHUNK_OVERLAP,
            length_function=len, # lambda x: sum([len(xi) for xi in x]),
        )
    
    async def process_file(self, file_path: str) -> List[LangchainDocument]:
        """
        Process a file and return chunks of text
        """
        file_extension = Path(file_path).suffix.lower()
        
        if file_extension == '.pdf':
            text = await self._extract_pdf_text(file_path)
        elif file_extension == '.docx':
            text = await self._extract_docx_text(file_path)
        elif file_extension == '.txt':
            text = await self._extract_txt_text(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_extension} - only one of [pdf, docx, txt] are supported")
        
        # Split text into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Convert chunks to Langchain documents
        documents = [
            LangchainDocument(
                page_content=chunk,
                metadata={
                    "source": file_path,
                    "chunk_index": i
                }
            )
            for i, chunk in enumerate(chunks)
        ]
        
        return documents
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """
        Extract text from PDF file using the PyPDF2 pacjage
        """
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """
        Extract text from DOCX file using the docx package
        """
        doc = Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    async def _extract_txt_text(self, file_path: str) -> str:
        """
        Extract text from TXT file with native python io
        """
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    async def save_processed_document(self, file_path: str, documents: List[LangchainDocument]):
        """
        Save processed document chunks to disk
        """
        # Create a directory for the document
        doc_name = Path(file_path).stem
        doc_dir = settings.PROCESSED_DATA_DIR / doc_name
        os.makedirs(doc_dir, exist_ok=True)
        
        # Save each chunk
        for i, doc in enumerate(documents):
            chunk_path = doc_dir / f"chunk_{i}.txt"
            with open(chunk_path, 'w', encoding='utf-8') as f:
                f.write(doc.page_content)
            
            # Save metadata
            metadata_path = doc_dir / f"chunk_{i}_metadata.json"
            with open(metadata_path, 'w', encoding='utf-8') as f:
                f.write(str(doc.metadata)) 